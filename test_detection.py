"""
测试脚本 - 使用模拟数据验证各检测方案
"""
from docx import Document
from pathlib import Path

from method_tfidf import detect_similarity_tfidf
from method_sequence import detect_similarity_sequence
from method_simhash import detect_similarity_simhash
from method_typo import (
    find_common_typos,
    find_common_unusual_phrases,
    find_common_punctuation_errors,
)


def create_test_docs():
    """创建测试用的 Word 文档"""
    test_dir = Path("test_docs")
    test_dir.mkdir(exist_ok=True)

    doc_a = Document()
    doc_a.add_paragraph("本项目位于某市某区，总建筑面积约为五万平方米，包括地上三十层和地下两层。")
    doc_a.add_paragraph("施工过程中应严格按照国家现行规范和标准执行，确保工程质量达到优良等级。")
    doc_a.add_paragraph("我公司拟投入项目经理一名、技术负责人一名、质量员两名、安全员两名。")
    doc_a.add_paragraph("本工程采用商品混凝土，强度等级为C30，抗渗等级为P8。")
    doc_a.add_paragraph("安全文明施工措施费按照招标文件要求计取，不得作为竟标优惠条件。")  # 错别字：竟→竞
    doc_a.save(test_dir / "doc_a.docx")

    doc_b = Document()
    doc_b.add_paragraph("本项目位于某市某区，总建筑面积约为五万平方米，包括地上三十层和地下两层。")  # 完全一致
    doc_b.add_paragraph("施工中应严格按照国家现行规范及标准执行，确保工程质量达到优良等级。")  # 高度相似
    doc_b.add_paragraph("我公司计划投入项目经理一名、技术负责人一名、质量员三名、安全员两名。")  # 部分相似
    doc_b.add_paragraph("本工程使用预拌混凝土，强度等级C30，抗渗等级P8，满足设计要求。")  # 中度相似
    doc_b.add_paragraph("安全文明施工措施费按照招标文件要求计取，不得作为竟标优惠条件。")  # 相同错别字
    doc_b.save(test_dir / "doc_b.docx")

    return str(test_dir / "doc_a.docx"), str(test_dir / "doc_b.docx")


def test_all_methods():
    """测试所有检测方案"""
    paragraphs_a = [
        "本项目位于某市某区，总建筑面积约为五万平方米，包括地上三十层和地下两层。",
        "施工过程中应严格按照国家现行规范和标准执行，确保工程质量达到优良等级。",
        "我公司拟投入项目经理一名、技术负责人一名、质量员两名、安全员两名。",
        "本工程采用商品混凝土，强度等级为C30，抗渗等级为P8。",
        "安全文明施工措施费按照招标文件要求计取，不得作为竟标优惠条件。",
    ]

    paragraphs_b = [
        "本项目位于某市某区，总建筑面积约为五万平方米，包括地上三十层和地下两层。",
        "施工中应严格按照国家现行规范及标准执行，确保工程质量达到优良等级。",
        "我公司计划投入项目经理一名、技术负责人一名、质量员三名、安全员两名。",
        "本工程使用预拌混凝土，强度等级C30，抗渗等级P8，满足设计要求。",
        "安全文明施工措施费按照招标文件要求计取，不得作为竟标优惠条件。",
    ]

    print("=" * 60)
    print("方案一: TF-IDF + 余弦相似度")
    print("=" * 60)
    results = detect_similarity_tfidf(paragraphs_a, paragraphs_b, threshold=0.5)
    for r in results:
        print(f"  相似度: {r['similarity']:.2%}")
        print(f"  A[{r['para_a_index']}]: {r['text_a'][:50]}...")
        print(f"  B[{r['para_b_index']}]: {r['text_b'][:50]}...")
        print()

    print("=" * 60)
    print("方案二: SequenceMatcher")
    print("=" * 60)
    results = detect_similarity_sequence(paragraphs_a, paragraphs_b, threshold=0.5)
    for r in results:
        print(f"  相似度: {r['similarity']:.2%}")
        print(f"  A[{r['para_a_index']}]: {r['text_a'][:50]}...")
        print(f"  B[{r['para_b_index']}]: {r['text_b'][:50]}...")
        if r.get("matching_blocks"):
            print(f"  公共子串: {r['matching_blocks'][:3]}")
        print()

    print("=" * 60)
    print("方案三: SimHash")
    print("=" * 60)
    results = detect_similarity_simhash(paragraphs_a, paragraphs_b, max_distance=15)
    for r in results:
        print(f"  相似度: {r['similarity']:.2%} (汉明距离: {r['hamming_distance']})")
        print(f"  A[{r['para_a_index']}]: {r['text_a'][:50]}...")
        print(f"  B[{r['para_b_index']}]: {r['text_b'][:50]}...")
        print()

    print("=" * 60)
    print("方案四: 共同错别字检测")
    print("=" * 60)
    typos = find_common_typos(paragraphs_a, paragraphs_b)
    print(f"  共同错别字: {len(typos)} 个")
    for t in typos:
        print(f"    - '{t['typo']}' (正确: '{t['correct_form']}')")
        print(f"      A上下文: {t['location_a']['context']}")
        print(f"      B上下文: {t['location_b']['context']}")

    phrases = find_common_unusual_phrases(paragraphs_a, paragraphs_b)
    print(f"\n  共同异常短语: {len(phrases)} 个")
    for p in phrases[:10]:
        print(f"    - '{p['phrase']}' (A:{p['count_in_a']}次, B:{p['count_in_b']}次)")


if __name__ == "__main__":
    test_all_methods()
