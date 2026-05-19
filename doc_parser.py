"""
Word 文档文本提取工具
"""
from docx import Document


def extract_paragraphs(docx_path: str) -> list[str]:
    """从 Word 文档中提取所有非空段落文本"""
    doc = Document(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def extract_table_cells(docx_path: str) -> list[str]:
    """从 Word 文档表格中提取所有非空单元格文本"""
    doc = Document(docx_path)
    cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    cells.append(text)
    return cells


def extract_all_text(docx_path: str) -> list[str]:
    """提取文档中所有文本（段落 + 表格）"""
    return extract_paragraphs(docx_path) + extract_table_cells(docx_path)
